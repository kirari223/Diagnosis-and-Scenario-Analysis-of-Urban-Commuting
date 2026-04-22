from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# --- styles ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph_with_links(doc, text):
    """Parse markdown links [text](url) and add as hyperlinks."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r'(\[.*?\]\(.*?\))', text)
    for part in parts:
        m = re.match(r'\[(.*?)\]\((.*?)\)', part)
        if m:
            link_text, url = m.group(1), m.group(2)
            add_hyperlink(p, link_text, url)
        else:
            if part:
                run = p.add_run(part)
                run.font.size = Pt(11)
    return p

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# ============================================================
# Title
# ============================================================
add_heading(doc, '文献综述参考：职住调整机制 & 情景分析范式', 1)
doc.add_paragraph('生成日期：2026-04-18').runs[0].font.size = Pt(10)

# ============================================================
# Section 1
# ============================================================
add_heading(doc, '一、职住调整机制（约1000-1500字）', 2)

add_paragraph_with_links(doc, '核心逻辑：这一节要回答"当通勤不满意时，人们如何调整？调整的异质性从何而来？"，为后面情景分析中引入"刚性"概念埋下伏笔。')

add_heading(doc, '1. 通勤不满意的触发与调整行为的理论基础（约200字）', 3)
add_paragraph_with_links(doc,
    '从经典的搜索理论（search theory）切入。van Ommeren et al. (1999) 在 [Job mobility, residential mobility and commuting](https://link.springer.com/article/10.1007/s001689900004) 中用搜索理论证明：换工作和搬家是同一个优化问题的两个解，通勤距离是触发搜索行为的核心变量。这给"刚性"概念提供了微观基础——刚性高的人，搜索成本高，调整行为被抑制。'
)

add_heading(doc, '2. 两条调整路径：搬家 vs 换工作（约400字）', 3)
add_paragraph_with_links(doc, '这是这一节的主体。文献上有清晰的分野：')

add_paragraph_with_links(doc,
    '搬家路径：Rouwendal & Meijer (2001) 等研究表明，居住迁移对通勤距离的响应是缓慢的，存在明显的"调整滞后"。双职工家庭（two-earner households）是典型的高刚性群体——[Spatial Moving Behavior of Two-Earner Households](https://www.researchgate.net/publication/229806140_Spatial_Moving_Behavior_of_Two-Earner_Households) 指出，两人的工作地约束相互牵制，使得搬家决策极为复杂，这直接对应模型中 θ 较高的情形。'
)
add_paragraph_with_links(doc,
    '换工作路径：职业流动性高的群体（如IT、知识密集型服务业）更倾向于通过换工作来优化通勤，而非搬家。这与情景二中IT行业低刚性的设定高度吻合。'
)
add_paragraph_with_links(doc,
    '工作地被动迁移：[Short distance, big impact](https://www.researchgate.net/publication/335039010_Short_distance_big_impact_The_effects_of_intra-city_workplace_relocation_on_staff_mobility_practices) 和 [University of Luxembourg Relocation](https://mdpi.com/2071-1050/12/18/7506) 研究了工作地被动迁移（即企业搬迁）对通勤行为的冲击，发现即使是短距离迁移也会显著重塑通勤格局——这正是情景二"产业空间强制转移"的现实原型。'
)

add_heading(doc, '3. 调整异质性：谁更刚性？（约400字）', 3)
add_paragraph_with_links(doc, '这是连接文献综述与方法框架的关键段落，要明确点出"刚性"的社会经济来源：')
add_paragraph_with_links(doc,
    '行业属性：制造业工人因工作地点固定（重资产、生产线绑定）、技能专用性强，调整能力弱；IT/知识经济从业者工作地点灵活，调整能力强。[The Effects of Employment Center Characteristics on Commuting Time](https://www.mdpi.com/2220-9964/14/3/116/htm) 中首尔的案例印证了制造业与金融/服务业在通勤时间上的系统性差异。'
)
add_paragraph_with_links(doc, '家庭结构：双职工、有子女家庭的调整刚性显著高于单身群体。')
add_paragraph_with_links(doc,
    '居家办公的特殊性：远程办公实质上是一种"不搬家、不换工作"的第三条路径，它通过解除空间约束来降低通勤刚性。[How working from home reshapes cities (PNAS 2024)](https://www.pnas.org/doi/10.1073/pnas.2408930121) 和 [Spatial Implications of Telecommuting](https://papers.ssrn.com/sol3/papers.cfm?abstract_id=3746555) 提供了宏观层面的证据。这里只需一笔带过，点明"居家办公是刚性的对立面"即可，不必展开。'
)

add_heading(doc, '4. 小结过渡（约100字）', 3)
add_paragraph_with_links(doc,
    '总结：调整行为的异质性本质上是通勤刚性的差异化表达，而刚性的来源是行业属性、家庭结构、住房市场摩擦等多重因素的叠加。这为后文引入UOT框架中的 θ 参数提供了现实依据。'
)

# ============================================================
# Section 2
# ============================================================
add_heading(doc, '二、情景分析范式（约1000-1500字）', 2)

add_paragraph_with_links(doc,
    '核心逻辑：这一节要完成一个"范式演进"的叙事——从城市规划的传统情景分析，到职住通勤领域的具体应用，再到最优传输理论作为新范式的引入。三个层次，层层递进。'
)

add_heading(doc, '1. 城市规划中的情景分析传统（约300字）', 3)
add_paragraph_with_links(doc,
    '[From forecasts to scenarios in strategic city-regional land-use and transportation planning](https://www.tandfonline.com/doi/full/10.1080/00343404.2022.2058699) 提供了一个很好的理论框架综述，可以用来说明情景分析从"预测"转向"探索性情景"的方法论演变。'
)
add_paragraph_with_links(doc,
    '核心要点：情景分析不是预测未来，而是通过构造反事实（counterfactual）来评估政策干预的边界效应。这个定性说明为后面的定量模型奠定方法论合法性。'
)

add_heading(doc, '2. 职住通勤领域的情景分析应用（约500字）', 3)
add_paragraph_with_links(doc, '这是这一节的主体，串联以下几条文献线索：')
add_paragraph_with_links(doc,
    '过剩通勤框架：Hamilton (1982) 的"wasteful commuting"开创了用最优化基准来衡量现实通勤效率的传统。[Extended Excess Commuting (Seoul)](https://journals.sagepub.com/doi/10.1080/00420980600945245) 和 [A big data approach to mitigating the MAUP](https://www.researchgate.net/publication/389708468_A_big_data_approach_to_mitigating_the_MAUP_in_measuring_excess_commuting) 是这一传统的延续。过剩通勤本质上就是一种隐性情景分析——"如果职住完全匹配，通勤会是什么样？"'
)
add_paragraph_with_links(doc,
    'Brotchie三角形：刘贤腾等(2018)引入的Brotchie三角形是一个优雅的情景框架，将城市通勤效率置于"最优"（最短通勤）、"最差"（最长通勤）和"现实"三个顶点之间，直观呈现政策干预的可能空间。这是一种几何化的情景分析，适合做横向城市比较。'
)
add_paragraph_with_links(doc,
    '空间相互作用模型情景：岳丽莹(2022)用单约束空间相互作用模型模拟通勤情景变化，代表了将引力模型/Wilson模型用于情景推演的技术路线。[Understanding commuting patterns and changes: Counterfactual analysis in a planning support framework](https://journals.sagepub.com/doi/abs/10.1177/2399808320924433) 的上海案例也属于这一路线，其反事实分析发现工作地布局对通勤的影响大于居住地布局，这与情景二的政策逻辑（调整就业地 D₀）直接呼应。'
)
add_paragraph_with_links(doc,
    '两种方法的共同局限：都是"硬约束"模型——要么完全服从规划，要么完全不服从，无法刻画现实中"部分服从、部分偏离"的弹性响应。这就是引入最优传输的动机。'
)

add_heading(doc, '3. 最优传输理论：从硬约束到软约束（约400字）', 3)
add_paragraph_with_links(doc,
    '经典最优传输（Monge-Kantorovich问题）：在给定供需约束下寻找最小成本匹配，对应的是"完全刚性"的极端情形——所有人都服从最优分配。这在数学上等价于过剩通勤框架中的"最优基准"。'
)
add_paragraph_with_links(doc,
    '非平衡最优传输（UOT）：通过引入KL散度惩罚项，允许供需约束被"软违反"，惩罚强度由 θ 控制。θ → ∞ 退化为硬约束（完全刚性），θ → 0 退化为完全弹性（无约束）。这在数学上为"刚性异质性"提供了统一的参数化表达。'
)
add_paragraph_with_links(doc,
    '与Wilson模型的关系：Wilson双约束模型在特定条件下是最优传输的一个特例（熵正则化最优传输），这建立了新旧方法之间的理论连续性，避免让读者觉得UOT是凭空引入的。'
)

add_heading(doc, '4. 小结过渡（约100字）', 3)
add_paragraph_with_links(doc,
    '总结：从Brotchie三角形到过剩通勤，再到UOT，情景分析的核心问题始终是"现实通勤与最优通勤之间的距离"，区别在于对"偏离"的建模方式越来越精细。UOT通过 θ 参数将这种偏离的行业异质性纳入统一框架，为情景二的产业空间重构分析提供了方法基础。'
)

# ============================================================
# Section 3: Supplement for Scenario 2
# ============================================================
add_heading(doc, '三、情景二文献支撑补充建议', 2)

add_paragraph_with_links(doc,
    '行业刚性的参数化依据：[The Effects of Employment Center Characteristics on Commuting Time](https://www.mdpi.com/2220-9964/14/3/116/htm) 中制造业 vs 服务业的通勤时间差异，可以作为 θ_mech > θ_it 的经验支撑，而不只是假设。'
)
add_paragraph_with_links(doc,
    '产业空间集聚的现实文献：[Short distance, big impact](https://www.researchgate.net/publication/335039010_Short_distance_big_impact_The_effects_of_intra-city_workplace_relocation_on_staff_mobility_practices) 和 [University of Luxembourg Relocation](https://mdpi.com/2071-1050/12/18/7506) 提供了工作地被动迁移的微观证据，可以用来论证"强制转移30%就业"这一冲击设定的现实合理性。'
)
add_paragraph_with_links(doc,
    '"一刀切规划"的批判：[Understanding commuting patterns: Counterfactual analysis](https://journals.sagepub.com/doi/abs/10.1177/2399808320924433) 的结论（工作地布局影响大于居住地）可以用来铺垫核心论点——忽视行业刚性异质性的产业空间规划会产生系统性偏差。'
)

# ============================================================
# Sources
# ============================================================
add_heading(doc, '参考文献链接汇总', 2)
sources = [
    ('Job mobility, residential mobility and commuting — van Ommeren et al. (1999)', 'https://link.springer.com/article/10.1007/s001689900004'),
    ('Spatial Moving Behavior of Two-Earner Households', 'https://www.researchgate.net/publication/229806140_Spatial_Moving_Behavior_of_Two-Earner_Households'),
    ('Short distance, big impact: intra-city workplace relocation', 'https://www.researchgate.net/publication/335039010_Short_distance_big_impact_The_effects_of_intra-city_workplace_relocation_on_staff_mobility_practices'),
    ('University of Luxembourg Relocation (Sustainability 2020)', 'https://mdpi.com/2071-1050/12/18/7506'),
    ('How working from home reshapes cities (PNAS 2024)', 'https://www.pnas.org/doi/10.1073/pnas.2408930121'),
    ('Spatial Implications of Telecommuting — Delventhal & Parkhomenko', 'https://papers.ssrn.com/sol3/papers.cfm?abstract_id=3746555'),
    ('From forecasts to scenarios in strategic city-regional LUT planning', 'https://www.tandfonline.com/doi/full/10.1080/00343404.2022.2058699'),
    ('Extended Excess Commuting: A Measure of the Jobs-Housing Imbalance in Seoul', 'https://journals.sagepub.com/doi/10.1080/00420980600945245'),
    ('A big data approach to mitigating the MAUP in measuring excess commuting', 'https://www.researchgate.net/publication/389708468_A_big_data_approach_to_mitigating_the_MAUP_in_measuring_excess_commuting'),
    ('Understanding commuting patterns and changes: Counterfactual analysis (Shanghai)', 'https://journals.sagepub.com/doi/abs/10.1177/2399808320924433'),
    ('The Effects of Employment Center Characteristics on Commuting Time (Seoul)', 'https://www.mdpi.com/2220-9964/14/3/116/htm'),
]
for label, url in sources:
    p = doc.add_paragraph(style='List Bullet')
    add_hyperlink(p, label, url)

doc.save(r'E:\00_Commute_Scenario_Research\docs\0418文献综述参考.docx')
print('saved')
